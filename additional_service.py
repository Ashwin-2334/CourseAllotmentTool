from flask import Blueprint, request, send_from_directory
import os
import pandas as pd
from docx import Document
from openpyxl import load_workbook

# Blueprint for additional service
additional_service = Blueprint('additional_service', __name__)

@additional_service.route('/process_additional', methods=['POST'])
def process_additional():
    print(f"Received files: {request.files.keys()}")  # Debugging line

    # Check for required files
    if 'preferencesExcel' not in request.files or 'mainExcel' not in request.files or 'schemaExcel' not in request.files:
        return "All three files (Preferences, Main Excel, and Schema Excel) are required!", 400

    # Retrieve files from the request
    preferences_file = request.files['preferencesExcel']
    main_file = request.files['mainExcel']
    schema_file = request.files['schemaExcel']

    # Save files to the uploads directory
    preferences_file_path = os.path.join('uploads', preferences_file.filename)
    main_file_path = os.path.join('uploads', main_file.filename)
    schema_file_path = os.path.join('uploads', schema_file.filename)

    preferences_file.save(preferences_file_path)
    main_file.save(main_file_path)
    schema_file.save(schema_file_path)

    # Pass all three paths to process_files
    output_file_path = process_files(preferences_file_path, main_file_path, schema_file_path)

    if not isinstance(output_file_path, str):
        return "Error processing files. File path invalid!", 500

    return f'''
    <h1>Files Processed Successfully!</h1>
    <a href="/download/{os.path.basename(output_file_path)}">Click here to download the processed file</a>
    '''

def process_files(preferences_excel_path, main_excel_path, schema_excel_path):
    """
    Process the preferences, main, and schema Excel files:
    - Assign unassigned courses based on preferences from the original Excel.
    - Generate a Word file containing:
      1. Course titles and assigned faculties for sections A, B, and C.
      2. A table for faculties still unassigned.
    """
    import pandas as pd
    from docx import Document
    import os

    print(f"Preferences Excel Path: {preferences_excel_path}")
    print(f"Main Excel Path: {main_excel_path}")
    print(f"Schema Excel Path: {schema_excel_path}")

    os.makedirs('processed', exist_ok=True)  # Ensure the processed folder exists

    try:
        # Load all Excel files
        preferences_data = pd.read_excel(preferences_excel_path)
        main_excel_data = pd.read_excel(main_excel_path)
        schema_data = pd.read_excel(schema_excel_path)

        # Normalize text for matching
        preferences_data = preferences_data.fillna("")
        main_excel_data["4th Semester"] = main_excel_data["4th Semester"].astype(str).str.strip()
        main_excel_data["6th Semester"] = main_excel_data["6th Semester"].astype(str).str.strip()
        schema_data["Course Title"] = schema_data["Course Title"].astype(str).str.strip()

        # Prepare sections
        sections = ['A', 'B', 'C']

        # Create a dictionary to store assignments
        course_assignments = {course: {"A": "No faculty assigned", "B": "No faculty assigned", "C": "No faculty assigned"} for course in schema_data["Course Title"].unique()}

        # Track assigned faculties
        assigned_faculty = set()

        # Assign faculties to sections based on 4th and 6th semester courses
        for course in schema_data["Course Title"].unique():
            faculties_for_course = main_excel_data[
                (main_excel_data["4th Semester"] == course) |
                (main_excel_data["6th Semester"] == course)
            ]
            for i, index in enumerate(faculties_for_course.index):
                if i < len(sections):  # Only assign A, B, C
                    course_assignments[course][sections[i]] = main_excel_data.at[index, "Faculty Name"]
                    assigned_faculty.add(main_excel_data.at[index, "Faculty Name"])

        # Identify unassigned faculties
        unassigned_faculties = main_excel_data[~main_excel_data["Faculty Name"].isin(assigned_faculty)]

        # Suggest unassigned courses based on preferences
        unassigned_courses = {
            course for course, sections in course_assignments.items()
            if "No faculty assigned" in sections.values()
        }

        for _, faculty_row in unassigned_faculties.iterrows():
            faculty_name = faculty_row["Faculty Name"]

            # Check preferences from the original preferences file
            faculty_preferences = preferences_data[
                preferences_data["Faculty Name"] == faculty_name
            ].iloc[0]

            suggested_course = "No courses available"
            for pref_num in range(1, 5):  # Assuming up to 4 preferences
                pref_col = f"Subject Preference {pref_num}"
                if pref_col in faculty_preferences and faculty_preferences[pref_col] in unassigned_courses:
                    suggested_course = faculty_preferences[pref_col]
                    break

            # Assign the suggested course dynamically
            if suggested_course in course_assignments and "No faculty assigned" in course_assignments[suggested_course].values():
                for section in sections:
                    if course_assignments[suggested_course][section] == "No faculty assigned":
                        course_assignments[suggested_course][section] = faculty_name
                        assigned_faculty.add(faculty_name)
                        break

        # Create a Word document
        doc = Document()
        doc.add_heading("Course Assignments by Section", level=1)

        # Add the first table for course assignments
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Set the header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Course Title"
        hdr_cells[1].text = "A Section"
        hdr_cells[2].text = "B Section"
        hdr_cells[3].text = "C Section"

        # Populate the table
        for course, sections in course_assignments.items():
            row_cells = table.add_row().cells
            row_cells[0].text = course
            row_cells[1].text = sections["A"]
            row_cells[2].text = sections["B"]
            row_cells[3].text = sections["C"]

        # Add a second table for unassigned faculties
        doc.add_heading("Unassigned Faculties", level=1)
        if unassigned_faculties.empty:
             doc.add_paragraph("All faculties have been assigned courses.")
        else:
            unassigned_table = doc.add_table(rows=1, cols=2)
            unassigned_table.style = 'Table Grid'

            # Set the header row
            hdr_cells = unassigned_table.rows[0].cells
            hdr_cells[0].text = "Faculty Name"
            hdr_cells[1].text = "Designation"

            # Populate the unassigned faculties table
            for _, row in unassigned_faculties.iterrows():
                row_cells = unassigned_table.add_row().cells
                row_cells[0].text = row["Faculty Name"]
                row_cells[1].text = row["Designation"]

        # Save the Word file
        output_file_path = os.path.join('processed', 'course_assignments.docx')
        doc.save(output_file_path)

        print(f"Processed File Saved at: {output_file_path}")
        return output_file_path

    except Exception as e:
        print(f"Error during processing: {e}")
        return f"Error: {str(e)}", 500
